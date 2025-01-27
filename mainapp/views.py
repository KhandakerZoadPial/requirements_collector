from django.shortcuts import render
from django.http import HttpResponse
from docx import Document



def collect_credentials(request):
    if request.method == 'POST':
        # Step 1: Parse data from the POST request
        figma_requirements = request.POST.getlist('figma_requirements')
        frontend_social_login = request.POST.getlist('frontend_social_login')
        backend_requirements = request.POST.getlist('backend_requirements')
        deployment_requirement = request.POST.get('deployment_requirements', None)  # Radio, so single value
        ai_requirements = request.POST.getlist('ai_requirements')

        print(figma_requirements, frontend_social_login, backend_requirements)

        # Step 2: Create a new Word document
        document = Document()
        document.add_heading('Client Credential Requirements', level=1)

        # Step 3: Add section-wise content

        # Figma Section
        if figma_requirements:
            document.add_heading('Figma (UI/UX)', level=2)
            for requirement in figma_requirements:
                document.add_paragraph(f"{requirement.replace('_', ' ').capitalize()}:")
                document.add_paragraph(" ")

        # Frontend Section
        if frontend_social_login:
            document.add_heading('Frontend (Social Login)', level=2)
            document.add_paragraph("Please create these dummy accounts and send us login information. Please turn off 2 factor authentiction")
            for login in frontend_social_login:
                platform = login.replace('_login', '').capitalize()
                document.add_paragraph(f"{platform} Email/Phone:")
                document.add_paragraph(f"{platform} Password:")
                document.add_paragraph(" ")

        # Backend Section
        if backend_requirements:
            document.add_heading('Backend', level=2)
            for requirement in backend_requirements:
                if requirement == 'mail_server_access':
                    document.add_paragraph("EMAIL_HOST: ")
                    document.add_paragraph("EMAIL_PORT: ")
                    document.add_paragraph("EMAIL_HOST_USER: ")
                    document.add_paragraph("EMAIL_HOST_PASSWORD: ")
                elif requirement == 'stripe_access':
                    document.add_paragraph("Stripe API Key:")
                    document.add_paragraph("Stripe Secret Key:")
                    document.add_paragraph("If it's subscription based payment, please create products in stripe, type=recurring")
                    document.add_paragraph("Price ID for Monthly: ")
                    document.add_paragraph("Price ID for Yearly: ")

                    document.add_paragraph("If you are unsure how to do this, send us the login credentials for stripe.")
                    document.add_paragraph("Stripe Email: ")
                    document.add_paragraph("Stripe Password: ")

                elif requirement == 'cloudinary_access':
                    document.add_paragraph("Cloudinary API Key:")
                    document.add_paragraph("Cloudinary API Secret:")
                document.add_paragraph(" ")

        # Deployment Section
        if deployment_requirement:
            document.add_heading('Deployment', level=2)
            deployment_platform = deployment_requirement.replace('_', ' ').capitalize()
            document.add_paragraph(f"Credentials for {deployment_platform}:")
            document.add_paragraph(f"{deployment_platform} email/username: ")
            document.add_paragraph(f"{deployment_platform} password: ")

        # AI Section
        if ai_requirements:
            document.add_heading('AI', level=2)
            for ai_key in ai_requirements:
                if ai_key == 'dataset_link':
                    document.add_paragraph("Link to Dataset:")
                elif ai_key == 'openai_key':
                    document.add_paragraph("OpenAI API Key:")
                elif ai_key == 'vapi_key':
                    document.add_paragraph("VAPI Key:")
                elif ai_key == 'deepgram_key':
                    document.add_paragraph("Deepgram API Key:")
                document.add_paragraph(" ")

        # Step 4: Save the document to response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=client_credentials.docx'
        document.save(response)
        return response

    return HttpResponse("Invalid request method", status=405)


# Create your views here.
def home(request):
    return render(request, 'home.html')